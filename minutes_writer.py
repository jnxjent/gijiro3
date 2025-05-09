import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_meiryo_font(run):
    """
    フォントを Meiryo に設定する関数。
    """
    run.font.name = 'Meiryo'
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

def write_minutes_section(word_file_path: str, output_file_path: str, replaced_transcription: str):
    """
    議事録をWordテンプレートに書き込む関数。

    ・テンプレート内の「■ 議事」段落を探し、その直後に replaced_transcription を挿入
    ・転記後、最後に「以上」を右寄せで追加
    ・全体のフォントを Meiryo に統一
    ・✅ 「■ 議事」の直後の不要な空白行を削除
    """
    doc = Document(word_file_path)

    # 「■ 議事」段落を探す
    marker_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "■ 議事" in paragraph.text:
            marker_index = i
            break
    if marker_index is None:
        print("[ERROR] 『■ 議事』の見出しが見つかりません。")
        return

    # ✅ 既存のブランク行を削除（「■ 議事」の直後）
    i = marker_index + 1
    while i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip():
        p = doc.paragraphs[i]
        p.clear()  # 空白段落を削除
        i += 1

    # ✅ 空白行を除外して議事録を挿入
    lines = [line for line in replaced_transcription.splitlines() if line.strip()]
    new_paragraphs = []
    for line in lines:
        new_para = doc.add_paragraph(line)
        new_paragraphs.append(new_para)

    # 「■ 議事」の直後に挿入
    doc.paragraphs[marker_index + 1:marker_index + 1] = new_paragraphs

    # 「以上」段落を追加
    end_para = doc.add_paragraph("以上")
    end_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 全体のフォントを Meiryo に統一
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_meiryo_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_meiryo_font(run)

    # 保存
    doc.save(output_file_path)
    print(f"[INFO] 議事録が保存されました: {output_file_path}")