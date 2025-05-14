from docx import Document
import unicodedata
import re

def normalize_text(text: str) -> str:
    """
    テキストを正規化し、表記ゆれを統一する。
    - NFKC 正規化 (全角・半角統一)
    - 末尾の「:」「：」を削除
    - 前後の空白を削除
    """
    return unicodedata.normalize("NFKC", text).strip().rstrip("：:")

def clean_repeated_labels(label: str, value: str) -> str:
    """
    `value` の冒頭に `label` が繰り返されている場合、それを削除する。
    - 記号 `:`, `：`, `-`, `～`, `ー` も考慮
    - "- label:" の形式も削除
    - ✅ 「1. - 議題:」を含んでいたら、その部分のみを削除
    """
    label = normalize_text(label)  # ラベルを正規化
    value = unicodedata.normalize("NFKC", value).strip()  # 値も正規化

    # `- label:` の形式や `label:` の繰り返しを削除
    pattern = rf"^[-\s]*{re.escape(label)}[\s:：\-～ー]*"
    cleaned_value = re.sub(pattern, '', value).strip()

    # ✅ 「1. - 議題:」を完全一致で削除
    cleaned_value = re.sub(r"1\. - 議題[:：]?", "", cleaned_value).strip()

    return cleaned_value


    # ✅ 「1. - 議題:」を部分一致で削除
    cleaned_value = re.sub(r"\d+\s*[-．]\s*議題\s*[:：]", "", cleaned_value).strip()

    return cleaned_value

def table_writer(word_file_path: str, output_file_path: str, extracted_info: dict):
    """
    WORDファイルの表から左列(0列)を抽出し、extracted_info のラベルと照合。
    一致する場合、右列(1列)に値を転記し、合致しない場合は空白のままにする。

    :param word_file_path: 読み込むWORDファイルのパス（テンプレート）
    :param output_file_path: 更新後のWORDファイルの保存先
    :param extracted_info: 生成AIが抽出した辞書データ {label: value}
    """
    doc = Document(word_file_path)

    # 📌 **ラベルの正規化とマッピング**
    label_map = {}  # {正規化ラベル: 元のラベル}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue  # 列が2つ未満の場合はスキップ
            raw_label = row.cells[0].text.strip()
            norm_label = normalize_text(raw_label)
            label_map[norm_label] = raw_label  # 正規化ラベル → 元ラベル

    print("[LOG] 正規化ラベルマップ:", label_map)  # デバッグ用ログ

    # 📌 **`extracted_info` のラベルも正規化**
    normalized_extracted_info = {normalize_text(k): v for k, v in extracted_info.items()}

    # 🔄 **テーブルの内容を `extracted_info` から取得し転記**
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue  

            raw_label = row.cells[0].text.strip()
            norm_label = normalize_text(raw_label)  # 正規化ラベルを取得

            # `normalized_extracted_info` 内のキーと照合
            if norm_label in normalized_extracted_info:
                raw_value = normalized_extracted_info[norm_label]  # 元の値
                cleaned_value = clean_repeated_labels(raw_label, raw_value)  # 繰り返し削除
                row.cells[1].text = str(cleaned_value)
                print(f"[LOG] {raw_label}: {cleaned_value} を転記")
            else:
                row.cells[1].text = ""  # 合致しない場合は空白
                print(f"[WARN] {raw_label} に対応するデータなし（空白のまま）")

    # 📄 **更新後のファイルを保存**
    doc.save(output_file_path)
    print("テーブルデータが保存されました:", output_file_path)
