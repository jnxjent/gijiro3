from docx import Document
import unicodedata
import re
import datetime

def normalize_text(text: str) -> str:
    """
    テキストを正規化し、表記ゆれを統一する。
    - NFKC 正規化 (全角・半角統一)
    - 末尾の「:」「：」を削除
    - 前後の空白を削除
    """
    return unicodedata.normalize("NFKC", text).strip().rstrip("：:")

def clean_repeated_labels(label: str, value: str) -> str:
    """
    `value` の冒頭に `label` が繰り返されている場合、それを削除する。
    - 記号 `:`, `：`, `-`, `～`, `ー` も考慮
    - "1. - 議題:" のような形式も削除
    """
    label_norm = normalize_text(label)
    # 文字列化＆NFKC正規化
    val = unicodedata.normalize("NFKC", str(value)).strip()

    # ① "- label:" や "label:" の繰り返しを削除
    pattern = rf"^[-\s]*{re.escape(label_norm)}[\s:：\-～ー]*"
    cleaned = re.sub(pattern, "", val).strip()

    # ② "数字. - ラベル:" の形式を削除
    cleaned = re.sub(
        rf"\d+\s*[\.．]\s*[-]?\s*{re.escape(label_norm)}\s*[:：]?", 
        "", 
        cleaned
    ).strip()

    return cleaned

def table_writer(word_file_path: str, output_file_path: str, extracted_info: dict):
    """
    WORDファイルの表から左列(0列)を抽出し、extracted_info のラベルと照合。
    一致する場合、右列(1列)に値を転記し、合致しない場合は空白のままにする。
    """

    doc = Document(word_file_path)

    # ── (1) テンプレート側のラベルを正規化してマッピング ─────────────────
    # 正規化ラベル → 元のセル文字列
    label_map = {}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            raw_label = row.cells[0].text.strip()
            norm_label = normalize_text(raw_label)
            label_map[norm_label] = raw_label

    print("[LOG] 正規化ラベルマップ:", label_map)

    # ── (2) extracted_info のキーも正規化 ─────────────────────────────────
    normalized_info = { normalize_text(k): v for k, v in extracted_info.items() }

    # ── (3) 各セルを走査して転記 ─────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            raw_label = row.cells[0].text.strip()
            norm_label = normalize_text(raw_label)

            if norm_label in normalized_info:
                raw_value = normalized_info[norm_label]

                # ── (A) 「出席者」は改行せずカンマ区切りに
                if norm_label == normalize_text("出席者"):
                    parts = re.split(r"[、,;\n\r]+", str(raw_value))
                    raw_value = ", ".join([p.strip() for p in parts if p.strip()])

                # ── (B) 「次回会議予定日時」は YYYY年MM月DD日 に整形
                if norm_label == normalize_text("次回会議予定日時"):
                    txt = str(raw_value)
                    year = datetime.datetime.now().year
                    def repl_date(m):
                        m_str, d_str = m.group(1), m.group(2)
                        return f"{year}年{int(m_str)}月{int(d_str)}日"
                    raw_value = re.sub(r"(\d{1,2})月\s*(\d{1,2})日", repl_date, txt)

                # ── (C) 重複ラベル削除
                cleaned = clean_repeated_labels(raw_label, raw_value)

                # ── (D) 行頭に「・」を追加（複数行対応）
                lines = cleaned.splitlines()
                bulleted = "\n".join(
                    f"・{ln}" if ln and not ln.startswith("・") else ln
                    for ln in lines
                )

                # 転記
                row.cells[1].text = bulleted
                print(f"[LOG] {raw_label}: {bulleted} を転記")

            else:
                # データなしは空白
                row.cells[1].text = ""
                print(f"[WARN] {raw_label} に対応するデータなし（空白に設定）")

    # ── (4) ファイル保存 ────────────────────────────────────────────────
    doc.save(output_file_path)
    print("テーブルデータが保存されました:", output_file_path)
